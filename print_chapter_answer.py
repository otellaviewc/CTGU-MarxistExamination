from typing import List, Dict
from docx import Document
import json

CHAPTER_COUNT = 5

chapter_numbers = ['绪论', '第一章', '第二章', '第三章', '第四章', '第五章', '第六章']


def get_problem_answer_text(problem: dict) -> str:
    problem_type = problem['type']
    if problem_type == 0:
        problem_answer = [problem['answer']]
    elif problem_type == 1:
        problem_answer = problem['answer']
    else:
        raise SystemExit

    return ''.join([chr(ord('A') + option) for option in problem_answer])


def add_chapter_answer(document, chapter: int, chapter_problems: List[dict]):
    # 统计题目数量
    problem_count = len(chapter_problems)

    # 添加标题
    document.add_heading(f'{chapter_numbers[chapter]}（共 {problem_count} 道题）', level=1)

    # 格式化
    w = len(f'{problem_count}')
    template = '{:0>%dd}—{:0>%dd}：' % (w, w)

    # 默认每章的答案按照题号排好序了
    for (idx, problem) in enumerate(chapter_problems):
        if idx % 5 == 0:
            paragraph = document.add_paragraph()
            paragraph.add_run(template.format(idx + 1, min(problem_count, idx + 5))).bold = True

        problem_answer_str = get_problem_answer_text(problem)
        paragraph.add_run(f'\t{problem_answer_str}')


def print_chapter_answer_docx(chapter_answer: Dict[int, List[dict]], save_path: str):
    # 生成文档
    doc = Document()
    doc.add_heading('马原每章答案', level=0)

    for i in range(CHAPTER_COUNT + 1):
        add_chapter_answer(doc, i, chapter_answer[i])

    # 保存生成的文档
    doc.save(save_path)


def print_chapter_answer_js_list(chapter_answer: Dict[int, List[dict]]):
    for i in range(CHAPTER_COUNT + 1):
        ls = ['']
        for problem in chapter_answer[i]:
            ls.append(get_problem_answer_text(problem))

        print(ls)


if __name__ == '__main__':
    # 读取答案
    answers_path = 'temp/answers.json'
    with open(answers_path, 'r') as answers_file:
        problems = json.load(answers_file)

    # 分割章节
    chapters: Dict[int, List[dict]] = {}
    for p in problems:
        chapters.setdefault(p['chapter'], []).append(p)

    print_chapter_answer_docx(chapters, "temp/马原每章答案.docx")
    # print_chapter_answer_js_list(chapters)
